import logging
import os
import re
import time
from datetime import datetime
from typing import Any, Dict, List, Optional

import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


logger = logging.getLogger(__name__)


class AppelsProjetFinder:
    """Service de recherche et extraction d'appels à projets.

    Étapes clés:
    - lecture d'un fichier Excel d'entrée
    - recherche via IA selon un 'sujet'
    - scraping d'URL et extraction assistée par IA
    - génération d'un document Word
    """

    def __init__(self, api_key: str, api_provider: str = "openai") -> None:
        key_clean = (api_key or "").strip()
        self.api_key = key_clean
        # Détection automatique des clés OpenRouter (sk-or-...)
        if key_clean.lower().startswith("sk-or-"):
            self.api_provider = "openrouter"
        else:
            self.api_provider = api_provider
        self.results: List[Dict[str, Any]] = []
        # Compteurs et diagnostics IA
        self.ai_calls: int = 0
        self.ai_success: int = 0
        self.heuristic_used: int = 0
        self.ai_errors: List[str] = []
        logger.info("Init AppelsProjetFinder provider=%s", self.api_provider)

    # -------------------- Lecture Excel --------------------
    def lire_fichier_excel(self, chemin_excel: str) -> Optional[pd.DataFrame]:
        try:
            logger.info("Lecture Excel: %s", chemin_excel)
            df = pd.read_excel(chemin_excel)
            logger.info("Excel lu: %d lignes", len(df))
            return df
        except Exception as exc:
            logger.exception("Erreur lecture Excel: %s", exc)
            return None

    # -------------------- Recherche IA --------------------
    def rechercher_avec_ia(self, sujet: str) -> List[Dict[str, Any]]:
        logger.info("Recherche IA sujet='%s' provider=%s", sujet, self.api_provider)
        if self.api_provider == "openai":
            return self._recherche_openai(sujet)
        if self.api_provider == "anthropic":
            return self._recherche_anthropic(sujet)
        if self.api_provider == "openrouter":
            return self._recherche_openrouter(sujet)
        logger.warning("Provider IA inconnu: %s", self.api_provider)
        return []

    def _call_openai_chat(self, messages: List[Dict[str, str]], model_legacy: str = "gpt-4o-mini", model_modern: str = "gpt-4o-mini") -> str:
        """Compatibilité OpenAI: utilise l'API 0.28 (ChatCompletion) ou 1.x (OpenAI client)."""
        try:
            import openai  # type: ignore
            version = getattr(openai, "__version__", "unknown")
            logger.info("OpenAI version=%s", version)

            # Cas SDK 0.28 (ancienne API)
            if hasattr(openai, "ChatCompletion"):
                openai.api_key = self.api_key
                logger.debug("OpenAI legacy ChatCompletion.create model=%s", model_legacy)
                response = openai.ChatCompletion.create(
                    model=model_legacy,
                    messages=messages,
                    temperature=0.3,
                )
                return response.choices[0].message.content  # type: ignore[attr-defined]

            # Cas SDK 1.x/2.x (nouvelle API)
            try:
                from openai import OpenAI  # type: ignore

                client = OpenAI(api_key=self.api_key)
                logger.debug("OpenAI modern chat.completions.create model=%s", model_modern)
                response = client.chat.completions.create(
                    model=model_modern,
                    messages=messages,
                    temperature=0.3,
                )
                return response.choices[0].message.content or ""
            except Exception as modern_exc:
                logger.exception("OpenAI modern client échec: %s", modern_exc)
                raise
        except Exception as exc:
            logger.exception("OpenAI indisponible: %s", exc)
            return ""

    def _recherche_openai(self, sujet: str) -> List[Dict[str, Any]]:
        prompt = (
            f"""Recherche les appels à projet actifs concernant: {sujet}

Pour chaque appel trouvé, extrais:
1. Titre de l'appel
2. Organisation responsable
3. Date de début des candidatures
4. Date de clôture des candidatures
5. Lien URL vers l'appel
6. Brève description (2-3 lignes)

Format de réponse en JSON:
[
  {{
    "titre": "...",
    "organisation": "...",
    "date_debut": "JJ/MM/AAAA",
    "date_cloture": "JJ/MM/AAAA",
    "url": "...",
    "description": "..."
  }}
]
"""
        )

        try:
            content = self._call_openai_chat(
                messages=[
                    {
                        "role": "system",
                        "content": "Tu es un assistant spécialisé dans la recherche d'appels à projet.",
                    },
                    {"role": "user", "content": prompt},
                ]
            )
            return self._parse_response(content)
        except Exception as exc:
            logger.exception("Erreur API OpenAI: %s", exc)
            return []

    def _recherche_anthropic(self, sujet: str) -> List[Dict[str, Any]]:
        try:
            import anthropic

            client = anthropic.Anthropic(api_key=self.api_key)
            prompt = (
                f"""Recherche les appels à projet actifs concernant: {sujet}

Pour chaque appel trouvé, extrais:
1. Titre de l'appel
2. Organisation responsable
3. Date de début des candidatures
4. Date de clôture des candidatures
5. Lien URL vers l'appel
6. Brève description (2-3 lignes)

Format de réponse en JSON:
[
  {{
    "titre": "...",
    "organisation": "...",
    "date_debut": "JJ/MM/AAAA",
    "date_cloture": "JJ/MM/AAAA",
    "url": "...",
    "description": "..."
  }}
]
"""
            )

            logger.debug("Appel Anthropic.messages.create")
            message = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}],
            )
            return self._parse_response(message.content[0].text)
        except Exception as exc:
            logger.exception("Erreur API Anthropic: %s", exc)
            return []

    # -------------------- OpenRouter --------------------
    def _call_openrouter_chat(self, messages: List[Dict[str, str]], model: Optional[str] = None) -> str:
        try:
            import json
            import requests

            api_key = self.api_key
            base_url = os.environ.get("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
            model_id = model or os.environ.get("OPENROUTER_MODEL", "openai/gpt-4o-mini")

            payload = {
                "model": model_id,
                "messages": messages,
                "temperature": 0.3,
            }

            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            }
            # En-têtes conseillés par OpenRouter (facultatifs mais utiles)
            referer = os.environ.get("OPENROUTER_REFERER")
            title = os.environ.get("OPENROUTER_TITLE", "ScrapingAP")
            if referer:
                headers["HTTP-Referer"] = referer
            if title:
                headers["X-Title"] = title

            resp = requests.post(f"{base_url}/chat/completions", headers=headers, data=json.dumps(payload), timeout=60)
            if resp.status_code >= 400:
                msg = f"OpenRouter HTTP {resp.status_code}: {resp.text[:300]}"
                logger.error(msg)
                self.ai_errors.append(msg)
                return ""
            data = resp.json()
            # Aligné sur la forme OpenAI
            content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
            if content:
                logger.info("IA ok provider=openrouter model=%s chars=%s", model_id, len(content))
            else:
                warn_msg = "OpenRouter: réponse vide"
                logger.warning(warn_msg)
                self.ai_errors.append(warn_msg)
            return content
        except Exception as exc:
            logger.exception("OpenRouter indisponible: %s", exc)
            self.ai_errors.append(f"OpenRouter exception: {exc}")
            return ""

    def _recherche_openrouter(self, sujet: str) -> List[Dict[str, Any]]:
        prompt = (
            f"""Recherche les appels à projet actifs concernant: {sujet}

Pour chaque appel trouvé, extrais:
1. Titre de l'appel
2. Organisation responsable
3. Date de début des candidatures
4. Date de clôture des candidatures
5. Lien URL vers l'appel
6. Brève description (2-3 lignes)

Format de réponse en JSON:
[
  {{
    "titre": "...",
    "organisation": "...",
    "date_debut": "JJ/MM/AAAA",
    "date_cloture": "JJ/MM/AAAA",
    "url": "...",
    "description": "..."
  }}
]
"""
        )

        content = self._call_openrouter_chat(
            messages=[
                {
                    "role": "system",
                    "content": "Tu es un assistant spécialisé dans la recherche d'appels à projet.",
                },
                {"role": "user", "content": prompt},
            ]
        )
        return self._parse_response(content)

    # -------------------- Scraping --------------------
    def scraper_url(self, url: str) -> Optional[Dict[str, Any]]:
        try:
            logger.info("Scraping URL: %s", url)
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                "Accept-Language": "fr-FR,fr;q=0.9,en-US;q=0.8,en;q=0.7",
            }
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code >= 400:
                # Beaucoup de sites renvoient 4xx/5xx mais ont tout de même du contenu exploitable
                logger.warning("HTTP %s sur %s — tentative d'extraction malgré tout", response.status_code, url)

            soup = BeautifulSoup(response.content or "", "html.parser")
            text_content = soup.get_text(separator=" ", strip=True)

            # 1) Tentative avec IA si disponible
            ia_result = self._extraire_info_avec_ia(text_content, url)
            if ia_result:
                return ia_result

            # 2) Fallback heuristique sans IA
            self.heuristic_used += 1
            logger.info("Heuristique utilisée url=%s (IA indisponible/échec)", url)
            return self._extraire_info_heuristique(soup, url)
        except Exception as exc:
            logger.exception("Erreur scraping %s", url)
            return None

    def _extraire_info_avec_ia(self, contenu: str, url: str) -> Optional[Dict[str, Any]]:
        logger.debug("Extraction IA depuis HTML url=%s contenu_len=%s", url, len(contenu))
        prompt = f"""Analyse ce contenu de page web et extrais les informations d'appel à projet:

{contenu[:4000]}

Extrais en JSON:
{{
  "titre": "...",
  "organisation": "...",
  "date_debut": "JJ/MM/AAAA",
  "date_cloture": "JJ/MM/AAAA",
  "description": "..."
}}

Si les informations ne sont pas disponibles, mets "Non spécifié".
"""

        try:
            if self.api_provider == "openai":
                self.ai_calls += 1
                content = self._call_openai_chat(messages=[{"role": "user", "content": prompt}])
                result = self._parse_response(content)
                if result:
                    result[0]["url"] = url
                    self.ai_success += 1
                    logger.info("IA ok provider=openai url=%s", url)
                    return result[0]
            elif self.api_provider == "openrouter":
                self.ai_calls += 1
                content = self._call_openrouter_chat(messages=[{"role": "user", "content": prompt}], model=os.environ.get("OPENROUTER_MODEL"))
                result = self._parse_response(content)
                if result:
                    result[0]["url"] = url
                    self.ai_success += 1
                    logger.info("IA ok provider=openrouter url=%s", url)
                    return result[0]
        except Exception as exc:
            logger.exception("Erreur extraction IA: %s", exc)
            self.ai_errors.append(str(exc))

        return None

    def _extraire_info_heuristique(self, soup: "BeautifulSoup", url: str) -> Optional[Dict[str, Any]]:
        """Extraction simple par heuristiques lorsque l'IA n'est pas disponible.

        Utilise title/meta/h1 et recherche des dates (FR/numériques) dans le texte.
        """
        try:
            def meta(prop: str, val: str) -> Optional[str]:
                tag = soup.find("meta", attrs={prop: val})
                if tag and tag.get("content"):
                    return str(tag.get("content")).strip()
                return None

            def first_text(selector_list: list[str]) -> Optional[str]:
                for sel in selector_list:
                    node = soup.select_one(sel)
                    if node and node.get_text(strip=True):
                        return node.get_text(strip=True)
                return None

            titre = (
                meta("property", "og:title")
                or meta("name", "twitter:title")
                or (soup.title.string.strip() if soup.title and soup.title.string else None)
                or first_text(["h1", "h2"])  # type: ignore[list-item]
            )
            description = (
                meta("property", "og:description")
                or meta("name", "description")
                or meta("name", "twitter:description")
            )
            organisation = meta("property", "og:site_name")

            full_text = soup.get_text(separator=" ", strip=True)

            # Dates en formats numériques ou texte FR
            mois = (
                "janvier|février|fevrier|mars|avril|mai|juin|juillet|août|aout|septembre|octobre|novembre|décembre|decembre"
            )
            date_num = r"\b\d{1,2}[\./-]\d{1,2}[\./-]\d{2,4}\b"
            date_txt = rf"\b\d{{1,2}}\s+(?:{mois})\s+\d{{4}}\b"
            date_pattern = rf"({date_num}|{date_txt})"

            re_opts = re.IGNORECASE | re.DOTALL
            m_cloture = re.search(rf"(cl[ôo]ture|deadline|limite).{{0,40}}{date_pattern}", full_text, re_opts)
            m_debut = re.search(rf"(d[ée]but|ouverture|start).{{0,40}}{date_pattern}", full_text, re_opts)

            date_cloture = None
            date_debut = None
            if m_cloture:
                date_cloture = m_cloture.group(2) if m_cloture.lastindex and m_cloture.lastindex >= 2 else m_cloture.group(0)
            if m_debut:
                date_debut = m_debut.group(2) if m_debut.lastindex and m_debut.lastindex >= 2 else m_debut.group(0)

            minimal = {
                "titre": titre or "Sans titre",
                "organisation": organisation or "Non spécifié",
                "date_debut": date_debut or "Non spécifiée",
                "date_cloture": date_cloture or "Non spécifiée",
                "url": url,
                "description": description or "",
            }

            # Si on n'a presque rien, retourne tout de même le titre + URL
            return minimal
        except Exception:
            logger.exception("Heuristique d'extraction échouée pour %s", url)
            return None

    # -------------------- Parsing --------------------
    def _parse_response(self, response_text: str) -> List[Dict[str, Any]]:
        try:
            import json

            logger.debug("Parsing réponse IA len=%s", len(response_text) if response_text else 0)
            json_match = re.search(r"\[.*\]", response_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            json_match = re.search(r"\{.*\}", response_text, re.DOTALL)
            if json_match:
                return [json.loads(json_match.group())]
        except Exception as exc:
            logger.exception("Erreur parsing JSON: %s", exc)
        return []

    # -------------------- Orchestration --------------------
    def traiter_fichier(self, chemin_excel: str) -> None:
        logger.info("Début traitement fichier: %s", chemin_excel)
        df = self.lire_fichier_excel(chemin_excel)
        if df is None:
            logger.error("Abandon: lecture Excel échouée")
            return

        # Normalise les noms de colonnes pour tolérer des synonymes
        original_to_lower = {c: c.strip().lower() for c in df.columns}
        df_proc = df.rename(columns=lambda c: original_to_lower[c])

        subject_candidates = [
            "sujet", "subject", "thème", "theme", "topic", "mots_clés", "mots-cles", "keywords"
        ]
        url_candidates = [
            "url", "lien", "link", "urls", "liens"
        ]

        subject_col = next((c for c in subject_candidates if c in df_proc.columns), None)
        url_col = next((c for c in url_candidates if c in df_proc.columns), None)

        if not subject_col and not url_col:
            logger.warning(
                "Aucune colonne reconnue. Attendu l'une de %s ou %s",
                subject_candidates,
                url_candidates,
            )

        self.results = []

        if subject_col:
            for idx, row in df_proc.iterrows():
                sujet = row[subject_col]
                if pd.notna(sujet) and str(sujet).strip():
                    logger.info("Recherche IA ligne=%s sujet=%s", idx, sujet)
                    self.results.extend(self.rechercher_avec_ia(str(sujet)))
                    time.sleep(1)

        if url_col:
            for idx, row in df_proc.iterrows():
                url = row[url_col]
                if pd.notna(url) and str(url).strip():
                    logger.info("Scraping ligne=%s url=%s", idx, url)
                    result = self.scraper_url(str(url))
                    if result:
                        self.results.append(result)
                    time.sleep(1)

        logger.info("Total résultats: %d", len(self.results))

    def generer_document_word(self, chemin_sortie: str = "appels_projet.docx") -> None:
        doc = Document()

        # Titre
        titre = doc.add_heading("Appels à Projet - Résultats de recherche", 0)
        titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_generation = doc.add_paragraph(
            f"Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        )
        date_generation.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()

        if not self.results:
            logger.warning("Aucun résultat à exporter - génération d'un document vide informatif")
            p = doc.add_paragraph()
            p.add_run("Aucun résultat trouvé pour les données fournies. ")
            p.add_run(
                "Vérifiez que votre fichier Excel contient une colonne 'sujet' ou 'url' (ou leurs synonymes: 'thème', 'link', 'lien', ...)."
            )
        else:
            for idx, appel in enumerate(self.results, 1):
                doc.add_heading(f"{idx}. {appel.get('titre', 'Sans titre')}", level=2)

                if appel.get("organisation"):
                    p = doc.add_paragraph()
                    p.add_run("Organisation: ").bold = True
                    p.add_run(str(appel["organisation"]))

                p = doc.add_paragraph()
                p.add_run("Date de début: ").bold = True
                p.add_run(str(appel.get("date_debut", "Non spécifiée")))

                p = doc.add_paragraph()
                p.add_run("Date de clôture: ").bold = True
                run = p.add_run(str(appel.get("date_cloture", "Non spécifiée")))
                if appel.get("date_cloture") and appel["date_cloture"] != "Non spécifiée":
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True

                if appel.get("url"):
                    p = doc.add_paragraph()
                    p.add_run("Lien: ").bold = True
                    p.add_run(str(appel["url"]))

                if appel.get("description"):
                    p = doc.add_paragraph()
                    p.add_run("Description: ").bold = True
                    p.add_run(str(appel["description"]))

                doc.add_paragraph("_" * 80)

        doc.save(chemin_sortie)
        logger.info("Document Word généré: %s", chemin_sortie)


