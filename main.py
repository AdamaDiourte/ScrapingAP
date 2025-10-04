"""
Outil de recherche automatique d'appels à projet
Auteur: Assistant IA
Date: Octobre 2025
"""

import os
from dotenv import load_dotenv
import re
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
from bs4 import BeautifulSoup
import time
import logging

logger = logging.getLogger(__name__)

# Charge les variables d'environnement depuis un fichier .env s'il existe
load_dotenv()


class AppelsProjetFinder:
    def __init__(self, api_key, api_provider="openai"):
        """
        Initialise l'outil de recherche d'appels à projet
        
        Args:
            api_key: Clé API pour l'IA (OpenAI, Anthropic, etc.)
            api_provider: "openai" ou "anthropic"
        """
        self.api_key = api_key
        self.api_provider = api_provider
        self.results = []
        logger.info(f"Initialisation AppelsProjetFinder provider={api_provider}")
        
    def lire_fichier_excel(self, chemin_excel):
        """
        Lit le fichier Excel contenant les sujets et/ou liens
        Format attendu: colonnes 'sujet' et/ou 'url'
        """
        try:
            logger.info(f"Lecture Excel: {chemin_excel}")
            df = pd.read_excel(chemin_excel)
            print(f"✓ Fichier Excel lu: {len(df)} entrées trouvées")
            logger.info(f"Excel lu: {len(df)} lignes")
            return df
        except Exception as e:
            print(f"✗ Erreur lecture Excel: {e}")
            logger.exception("Erreur lecture Excel")
            return None
    
    def rechercher_avec_ia(self, sujet):
        """
        Utilise l'API IA pour rechercher des appels à projet
        """
        logger.info(f"Recherche IA pour sujet='{sujet}' via provider={self.api_provider}")
        if self.api_provider == "openai":
            return self._recherche_openai(sujet)
        elif self.api_provider == "anthropic":
            return self._recherche_anthropic(sujet)
    
    def _recherche_openai(self, sujet):
        """Recherche via OpenAI API"""
        try:
            import openai
            openai.api_key = self.api_key
            
            prompt = f"""Recherche les appels à projet actifs concernant: {sujet}

Pour chaque appel trouvé, extrais:
1. Titre de l'appel
2. Organisation responsable
3. Date de début des candidatures
4. Date de clôture des candidatures
5. Lien URL vers l'appel
6. Brève description (2-3 lignes)

Format de réponse en JSON:
[
  {{
    "titre": "...",
    "organisation": "...",
    "date_debut": "JJ/MM/AAAA",
    "date_cloture": "JJ/MM/AAAA",
    "url": "...",
    "description": "..."
  }}
]
"""
            
            logger.debug("Appel OpenAI.ChatCompletion.create")
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Tu es un assistant spécialisé dans la recherche d'appels à projet."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            return self._parse_response(response.choices[0].message.content)
            
        except Exception as e:
            print(f"✗ Erreur API OpenAI: {e}")
            logger.exception("Erreur API OpenAI")
            return []
    
    def _recherche_anthropic(self, sujet):
        """Recherche via Anthropic Claude API"""
        try:
            import anthropic
            
            client = anthropic.Anthropic(api_key=self.api_key)
            
            prompt = f"""Recherche les appels à projet actifs concernant: {sujet}

Pour chaque appel trouvé, extrais:
1. Titre de l'appel
2. Organisation responsable
3. Date de début des candidatures
4. Date de clôture des candidatures
5. Lien URL vers l'appel
6. Brève description (2-3 lignes)

Format de réponse en JSON:
[
  {{
    "titre": "...",
    "organisation": "...",
    "date_debut": "JJ/MM/AAAA",
    "date_cloture": "JJ/MM/AAAA",
    "url": "...",
    "description": "..."
  }}
]
"""
            
            logger.debug("Appel Anthropic.messages.create")
            message = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=4096,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            return self._parse_response(message.content[0].text)
            
        except Exception as e:
            print(f"✗ Erreur API Anthropic: {e}")
            logger.exception("Erreur API Anthropic")
            return []
    
    def scraper_url(self, url):
        """
        Scrape une URL spécifique pour extraire les infos d'appel à projet
        """
        try:
            logger.info(f"Scraping URL: {url}")
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')
            text_content = soup.get_text(separator=' ', strip=True)

            # Utilise l'IA pour extraire les informations structurées
            return self._extraire_info_avec_ia(text_content, url)

        except Exception as e:
            print(f"✗ Erreur scraping {url}: {e}")
            logger.exception(f"Erreur scraping {url}")
            return None
    
    def _extraire_info_avec_ia(self, contenu, url):
        """Extrait les informations d'un contenu HTML avec l'IA"""
        logger.debug(f"Extraction IA depuis HTML url={url} contenu_len={len(contenu)}")
        prompt = f"""Analyse ce contenu de page web et extrais les informations d'appel à projet:

{contenu[:4000]}

Extrais en JSON:
{{
  "titre": "...",
  "organisation": "...",
  "date_debut": "JJ/MM/AAAA",
  "date_cloture": "JJ/MM/AAAA",
  "description": "..."
}}

Si les informations ne sont pas disponibles, mets "Non spécifié".
"""
        
        if self.api_provider == "openai":
            try:
                import openai
                openai.api_key = self.api_key
                
                logger.debug("Appel OpenAI.ChatCompletion.create pour extraction")
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3
                )
                
                result = self._parse_response(response.choices[0].message.content)
                if result and len(result) > 0:
                    result[0]['url'] = url
                    return result[0]
                    
            except Exception as e:
                print(f"✗ Erreur extraction IA: {e}")
                logger.exception("Erreur extraction IA")
        
        return None
    
    def _parse_response(self, response_text):
        """Parse la réponse JSON de l'IA"""
        try:
            import json
            logger.debug(f"Parsing réponse IA len={len(response_text) if response_text else 0}")
            # Cherche le JSON dans la réponse
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    return [json.loads(json_match.group())]
        except Exception as e:
            print(f"✗ Erreur parsing JSON: {e}")
            logger.exception("Erreur parsing JSON")
        return []
    
    def traiter_fichier(self, chemin_excel):
        """
        Traite le fichier Excel et collecte tous les résultats
        """
        logger.info(f"Début traitement fichier: {chemin_excel}")
        df = self.lire_fichier_excel(chemin_excel)
        if df is None:
            logger.error("Abandon: lecture Excel a échoué")
            return
        
        self.results = []
        
        # Traite les sujets
        if 'sujet' in df.columns:
            for idx, row in df.iterrows():
                sujet = row['sujet']
                if pd.notna(sujet):
                    print(f"\n🔍 Recherche pour: {sujet}")
                    logger.info(f"Recherche IA sujet ligne={idx}: {sujet}")
                    resultats = self.rechercher_avec_ia(sujet)
                    self.results.extend(resultats)
                    time.sleep(1)  # Pause pour éviter rate limiting
        
        # Traite les URLs spécifiques
        if 'url' in df.columns:
            for idx, row in df.iterrows():
                url = row['url']
                if pd.notna(url):
                    print(f"\n🌐 Scraping: {url}")
                    logger.info(f"Scraping URL ligne={idx}: {url}")
                    resultat = self.scraper_url(url)
                    if resultat:
                        self.results.append(resultat)
                    time.sleep(1)
        
        print(f"\n✓ Total: {len(self.results)} appels à projet trouvés")
        logger.info(f"Total résultats: {len(self.results)}")
    
    def generer_document_word(self, chemin_sortie="appels_projet.docx"):
        """
        Génère un document Word avec les résultats
        """
        if not self.results:
            print("✗ Aucun résultat à exporter")
            logger.warning("Aucun résultat à exporter")
            return
        
        doc = Document()
        
        # Titre
        titre = doc.add_heading('Appels à Projet - Résultats de recherche', 0)
        titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Date de génération
        date_generation = doc.add_paragraph(f"Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        date_generation.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()  # Espace
        
        # Résultats
        for idx, appel in enumerate(self.results, 1):
            # Titre de l'appel
            heading = doc.add_heading(f"{idx}. {appel.get('titre', 'Sans titre')}", level=2)
            
            # Organisation
            if appel.get('organisation'):
                p = doc.add_paragraph()
                p.add_run('Organisation: ').bold = True
                p.add_run(appel['organisation'])
            
            # Dates
            p = doc.add_paragraph()
            p.add_run('Date de début: ').bold = True
            p.add_run(appel.get('date_debut', 'Non spécifiée'))
            
            p = doc.add_paragraph()
            p.add_run('Date de clôture: ').bold = True
            run = p.add_run(appel.get('date_cloture', 'Non spécifiée'))
            # Colore en rouge si la date est proche
            if appel.get('date_cloture') and appel['date_cloture'] != 'Non spécifiée':
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
            
            # Lien
            if appel.get('url'):
                p = doc.add_paragraph()
                p.add_run('Lien: ').bold = True
                p.add_run(appel['url'])
            
            # Description
            if appel.get('description'):
                p = doc.add_paragraph()
                p.add_run('Description: ').bold = True
                p.add_run(appel['description'])
            
            doc.add_paragraph('_' * 80)  # Séparateur
        
        # Sauvegarde
        doc.save(chemin_sortie)
        print(f"\n✓ Document Word généré: {chemin_sortie}")
        logger.info(f"Document Word généré: {chemin_sortie}")


def main():
    """
    Fonction principale - Exemple d'utilisation
    """
    # Configuration via variables d'environnement (.env supporté)
    # Priorité: AP_FINDER_API_KEY > OPENAI_API_KEY > ANTHROPIC_API_KEY
    API_KEY = (
        os.getenv("AP_FINDER_API_KEY")
        or os.getenv("OPENAI_API_KEY")
        or os.getenv("ANTHROPIC_API_KEY")
    )
    API_PROVIDER = os.getenv("API_PROVIDER", "openai")  # ou "anthropic"
    
    FICHIER_EXCEL = "appels_projet_recherche.xlsx"
    FICHIER_SORTIE = "resultats_appels_projet.docx"
    
    # Vérification
    if not API_KEY:
        print("⚠️  ATTENTION: aucune clé API n'a été trouvée.")
        print("Définissez l'une des variables d'environnement suivantes: AP_FINDER_API_KEY, OPENAI_API_KEY ou ANTHROPIC_API_KEY.")
        print("Vous pouvez créer un fichier .env à la racine avec: AP_FINDER_API_KEY=... et éventuellement API_PROVIDER=openai|anthropic")
        return
    
    if not os.path.exists(FICHIER_EXCEL):
        print(f"⚠️  Le fichier {FICHIER_EXCEL} n'existe pas!")
        print("\nCréez un fichier Excel avec les colonnes:")
        print("  - 'sujet': sujets de recherche")
        print("  - 'url': URLs spécifiques à analyser")
        return
    
    # Exécution
    print("=" * 60)
    print("OUTIL DE RECHERCHE D'APPELS À PROJET")
    print("=" * 60)
    
    finder = AppelsProjetFinder(api_key=API_KEY, api_provider=API_PROVIDER)
    finder.traiter_fichier(FICHIER_EXCEL)
    finder.generer_document_word(FICHIER_SORTIE)
    
    print("\n✅ Traitement terminé!")


if __name__ == "__main__":
    main()